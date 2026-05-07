"""DOCX document parser using python-docx."""

from pathlib import Path


def parse_docx(file_path: Path) -> str:
    """Extract text from a DOCX file.

    Args:
        file_path: Path to the DOCX file

    Returns:
        Extracted text content
    """
    from docx import Document

    doc = Document(str(file_path))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                paragraphs.append(row_text)

    return "\n".join(paragraphs)
